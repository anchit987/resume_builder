import pdfplumber
import docx
import re
from typing import Dict, List, Tuple
import logging

logger = logging.getLogger(__name__)

def clean_text(text: str) -> str:
    """Simple but robust text cleaning."""
    if not text:
        return ""
    
    try:
        # Basic cleanup - remove excessive whitespace
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Remove excessive spaces
            line = ' '.join(line.split())
            if line.strip():
                cleaned_lines.append(line.strip())
        
        # Join with single newlines and reduce multiple empty lines
        result = '\n'.join(cleaned_lines)
        
        # Remove multiple consecutive newlines
        while '\n\n\n' in result:
            result = result.replace('\n\n\n', '\n\n')
            
        return result.strip()
        
    except Exception as e:
        logger.warning(f"Text cleaning failed: {e}")
        return text.strip()

def extract_contact_info(text: str) -> Dict[str, List[str]]:
    """Extract contact information with error handling."""
    contact_info = {}
    
    try:
        # Email - simple pattern
        emails = re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
        if emails:
            contact_info['emails'] = list(set(emails))  # Remove duplicates
        
        # Phone numbers - basic patterns
        phones = []
        phone_patterns = [
            r'\b\d{10}\b',  # 10 digit number
            r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',  # XXX-XXX-XXXX variations
            r'\(\d{3}\)\s*\d{3}[-.\s]?\d{4}',  # (XXX) XXX-XXXX
        ]
        
        for pattern in phone_patterns:
            matches = re.findall(pattern, text)
            phones.extend(matches)
        
        if phones:
            # Clean phone numbers
            cleaned_phones = []
            for phone in phones:
                # Remove all non-digits
                phone_digits = re.sub(r'\D', '', phone)
                if len(phone_digits) == 10:
                    cleaned_phones.append(phone.strip())
            contact_info['phones'] = list(set(cleaned_phones))
        
        # LinkedIn
        linkedin_matches = re.findall(r'linkedin\.com/in/[^\s]+', text, re.IGNORECASE)
        if linkedin_matches:
            contact_info['linkedin'] = [match.strip() for match in linkedin_matches]
        
        # GitHub
        github_matches = re.findall(r'github\.com/[^\s]+', text, re.IGNORECASE)
        if github_matches:
            contact_info['github'] = [match.strip() for match in github_matches]
            
    except Exception as e:
        logger.warning(f"Contact extraction failed: {e}")
        contact_info['extraction_error'] = str(e)
    
    return contact_info

def parse_pdf(file_path: str) -> str:
    """Simple but robust PDF parsing."""
    try:
        logger.info(f"Starting PDF parsing: {file_path}")
        full_text = ""
        
        with pdfplumber.open(file_path) as pdf:
            logger.info(f"PDF has {len(pdf.pages)} pages")
            
            for page_num, page in enumerate(pdf.pages, 1):
                logger.info(f"Processing page {page_num}")
                
                # Method 1: Simple text extraction
                page_text = page.extract_text()
                
                if not page_text or len(page_text.strip()) < 20:
                    logger.info(f"Page {page_num}: Trying alternative extraction")
                    # Method 2: With different settings
                    try:
                        page_text = page.extract_text(
                            x_tolerance=2,
                            y_tolerance=2
                        )
                    except Exception as e:
                        logger.warning(f"Alternative extraction failed: {e}")
                        page_text = ""
                
                if page_text:
                    full_text += f"\n=== PAGE {page_num} ===\n"
                    full_text += page_text + "\n"
                    logger.info(f"Page {page_num}: Extracted {len(page_text)} characters")
                else:
                    logger.warning(f"Page {page_num}: No text extracted")
                
                # Try to extract tables if main text is sparse
                if len(page_text.strip()) < 50:
                    try:
                        tables = page.find_tables()
                        for table in tables:
                            table_data = table.extract()
                            if table_data:
                                table_text = "\n".join([" | ".join([str(cell) if cell else "" for cell in row]) for row in table_data])
                                full_text += f"\n[TABLE DATA]\n{table_text}\n"
                                logger.info(f"Page {page_num}: Extracted table with {len(table_text)} characters")
                    except Exception as e:
                        logger.warning(f"Table extraction failed on page {page_num}: {e}")
        
        # Clean the extracted text
        cleaned_text = clean_text(full_text)
        
        # Extract contact info for validation
        contact_info = extract_contact_info(cleaned_text)
        logger.info(f"Contact info found: {list(contact_info.keys())}")
        
        if len(cleaned_text.strip()) < 100:
            logger.error("Very little text extracted - possible parsing issue")
            return f"[EXTRACTION WARNING: Only {len(cleaned_text)} characters extracted]\n\n{cleaned_text}"
        
        logger.info(f"Successfully extracted {len(cleaned_text)} characters from PDF")
        return cleaned_text
        
    except Exception as e:
        logger.error(f"PDF parsing failed: {str(e)}")
        raise Exception(f"PDF parsing error: {str(e)}")

def parse_docx(file_path: str) -> str:
    """Simple but robust DOCX parsing."""
    try:
        logger.info(f"Starting DOCX parsing: {file_path}")
        
        doc = docx.Document(file_path)
        full_text = ""
        
        # Extract paragraphs
        paragraph_count = 0
        for para in doc.paragraphs:
            if para.text.strip():
                full_text += para.text + "\n"
                paragraph_count += 1
        
        logger.info(f"Extracted {paragraph_count} paragraphs")
        
        # Extract tables
        table_count = 0
        for table in doc.tables:
            full_text += "\n[TABLE]\n"
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    full_text += row_text + "\n"
            full_text += "[/TABLE]\n\n"
            table_count += 1
        
        logger.info(f"Extracted {table_count} tables")
        
        # Clean the text
        cleaned_text = clean_text(full_text)
        
        if len(cleaned_text.strip()) < 50:
            logger.error("Very little text extracted from DOCX")
            return f"[EXTRACTION WARNING: Only {len(cleaned_text)} characters extracted]\n\n{cleaned_text}"
        
        logger.info(f"Successfully extracted {len(cleaned_text)} characters from DOCX")
        return cleaned_text
        
    except Exception as e:
        logger.error(f"DOCX parsing failed: {str(e)}")
        raise Exception(f"DOCX parsing error: {str(e)}")

# Test function
def test_parser(file_path: str):
    """Test the parser with a file."""
    try:
        ext = file_path.lower().split('.')[-1]
        if ext == 'pdf':
            result = parse_pdf(file_path)
        elif ext == 'docx':
            result = parse_docx(file_path)
        else:
            return "Unsupported file type"
        
        print(f"Extracted {len(result)} characters")
        print("First 500 characters:")
        print(result[:500])
        print("\n" + "="*50)
        print("Contact info:")
        contact_info = extract_contact_info(result)
        for key, value in contact_info.items():
            print(f"{key}: {value}")
            
        return result
        
    except Exception as e:
        print(f"Test failed: {e}")
        return None